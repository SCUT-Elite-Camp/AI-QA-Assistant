"""生成 Agent RL 论文综述测试文档"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("Agent 强化学习前沿论文综述", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "本文档系统整理了近年 Agent 强化学习领域的重要研究论文，"
    "涵盖推理增强、工具使用、自我改进与多智能体协作四大方向。"
)

# ── 一、推理增强 RL ──
doc.add_heading("一、推理增强 RL", 1)

doc.add_heading("1.1 DeepSeek-R1: Incentivizing Reasoning Capability in LLMs via RL", 2)
doc.add_paragraph("作者：DeepSeek-AI (2025)")
doc.add_paragraph(
    "DeepSeek-R1 首次在大规模语言模型上纯粹通过强化学习激发推理能力，免去传统的监督微调步骤。"
    "研究团队基于 DeepSeek-V3-Base 模型，直接应用 GRPO（Group Relative Policy Optimization）算法进行 RL 训练。"
    "训练过程中模型自发涌现了「顿悟时刻」（Aha Moment）——模型学会了在推理中途停下来重新审视已知条件、"
    "纠正先前错误、尝试替代方案等复杂推理行为，而这些行为完全未在奖励函数中显式编码。"
)
doc.add_paragraph(
    "技术要点：\n"
    "• 奖励设计：采用基于规则的可验证奖励（数学答案正确性 + LeetCode 测试用例通过率），避免 reward hacking\n"
    "• 训练模板：极简 prompt 模板，仅要求模型将推理过程置于 think 标签内，最终答案置于 answer 标签内\n"
    "• 冷启动数据：R1 最终版加入了数千条高质量冷启动 SFT 数据以改善输出可读性\n"
    "• 蒸馏：将 R1 的推理能力蒸馏到 1.5B~70B 的 Qwen/Llama 小模型上，小模型在数学推理基准上超越 GPT-4o"
)

doc.add_heading("1.2 STILL-ALIVE: Learning to Reason via Stepwise Internal Reward", 2)
doc.add_paragraph("作者：Zihan Wang et al. (2025)")
doc.add_paragraph(
    "该工作提出了一种名为 STILL-ALIVE（Stepwise In-Context Learning with LLM-Amortized Value Estimation）的方法，"
    "通过 LLM 自身作为价值估计器，为推理链中的每个中间步骤分配细粒度奖励信号。传统 RL 方法通常只在最终结果处给予稀疏奖励，"
    "而 STILL-ALIVE 利用 LLM 对每个推理步骤进行「前瞻性」价值评估——即判断当前步骤对最终答案的贡献程度。\n\n"
    "核心贡献：\n"
    "• 步级内在奖励（Stepwise Intrinsic Reward）：无需人工标注，LLM 自动评估中间步骤质量\n"
    "• 摊销价值估计：利用上下文学习能力进行价值迭代，大幅降低训练成本\n"
    "• 实验证明在 GSM8K、MATH 等数学推理任务上显著优于仅使用最终奖励的 baseline"
)

doc.add_heading("1.3 DeepEnlighten: Self-Improving Agent through Cyclic RL", 2)
doc.add_paragraph("作者：Meta AI (2025)")
doc.add_paragraph(
    "DeepEnlighten 提出了一种循环式 RL 框架来实现 Agent 的持续自我进化。其核心思想是将 Agent 的使用过程本身转化为训练信号——"
    "Agent 在执行任务时会产生成功/失败的轨迹，这些轨迹被自动收集、过滤并用于新一轮的 RL 训练，形成「训练→部署→收集→再训练」的闭环。\n\n"
    "架构设计：\n"
    "• Judge Model：自动评估轨迹质量，判断子任务完成度，筛选高质量训练样本\n"
    "• Cyclic Training Loop：每轮从部署中筛选 top-20% 轨迹混入训练集，逐步提升 Agent 能力边界\n"
    "• 在 SWE-Bench、WebArena 等真实环境基准上实现 SOTA，且迭代过程中性能单调提升"
)

# ── 二、工具使用与代码智能体 ──
doc.add_heading("二、工具使用与代码智能体", 1)

doc.add_heading("2.1 SWE-Agent: Agent-Computer Interfaces for Automated Software Engineering", 2)
doc.add_paragraph("作者：John Yang, Carlos E. Jimenez et al. (Princeton, 2024)")
doc.add_paragraph(
    "SWE-Agent 是首个在 SWE-Bench 上取得突破性成绩的 LM 编程智能体。它提出了一套精心设计的 Agent-Computer Interface（ACI）——"
    "即 Agent 与计算机交互的命令格式和约束规范，使得 LLM 能够像人类开发者一样在终端中浏览仓库、编辑文件、执行测试。\n\n"
    "关键设计：\n"
    "• ACI 原语：设计了 bash（执行命令）、edit（精确编辑文件）、scroll（浏览长文件）等专用命令\n"
    "• 自定义解析器：将 LLM 输出解析为可执行操作，而非自由格式的代码生成\n"
    "• 在 SWE-Bench Lite 上达到 18% 解决率（GPT-4），远超市面上其他方法\n"
    "• Linter 集成：在编辑操作后自动运行 linter 检查语法错误，发现错误则回退并重试"
)

doc.add_heading("2.2 OpenHands: A Platform for Code AI Agents", 2)
doc.add_paragraph("作者：Xingyao Wang et al. (2024)")
doc.add_paragraph(
    "OpenHands（原 OpenDevin）是一个开源的代码 AI Agent 平台，提供标准化的 Agent 开发与评测框架。"
    "该平台抽象出了一套通用的 Agent 行为接口（AgentLoop），支持多种 LLM 后端、多种沙箱环境、多种评估基准的即插即用。\n\n"
    "平台特性：\n"
    "• 事件流架构：Agent 的所有行为（观察、思考、行动）均作为事件在流中传播，便于监控与调试\n"
    "• 技能库：预置了文件编辑、终端执行、网页浏览、Git 操作等常用技能\n"
    "• 多 Agent 协作：支持 Delegator Agent 将子任务指派给专用子 Agent\n"
    "• 沙箱隔离：每个 Agent 实例运行在独立的 Docker 容器中，确保安全性"
)

doc.add_heading("2.3 Cline: Autonomous Coding Agent with Human-in-the-Loop", 2)
doc.add_paragraph("作者：Cline Community (2025)")
doc.add_paragraph(
    "Cline 是面向实际开发场景的自主编程 Agent，采用「规划-执行-验证」三步循环架构，"
    "在每步操作前向用户展示即将执行的命令和文件修改，支持一键批准或拒绝，实现了高效的人机协作模式。\n\n"
    "技术特点：\n"
    "• 分层规划：先制定高层计划（Plan Mode），用户确认后再逐步执行\n"
    "• 工具链整合：深度集成 VS Code / JetBrains IDE、终端、文件系统、浏览器等开发工具\n"
    "• MCP 协议：支持 Model Context Protocol 扩展，可接入自定义工具和外部 API\n"
    "• 安全护栏：文件修改前自动生成备份、执行高风险命令前额外提示确认"
)

# ── 三、多 Agent 协作与博弈 ──
doc.add_heading("三、多 Agent 协作与博弈", 1)

doc.add_heading("3.1 ChatDev: Communicative Agents for Software Development", 2)
doc.add_paragraph("作者：Chen Qian et al. (Tsinghua University, 2023)")
doc.add_paragraph(
    "ChatDev 开创性地将软件开发流程形式化为多智能体对话——不同角色（CEO、CTO、程序员、测试员、设计师）"
    "的 LLM Agent 通过自然语言聊天协作完成软件项目的全生命周期开发。它创新性地引入了「闲聊」（Chitchat）机制来促进信息交换，"
    "以及「自我反思」（Self-Reflection）机制让 Agent 在产生分歧时回溯并修正自己的输出。\n\n"
    "核心机制：\n"
    "• 角色专业化：每个 Agent 拥有独立的系统提示和职责范围（需求分析、架构设计、编码、测试、文档）\n"
    "• 阶段式工作流：将开发过程分为 Design → Coding → Testing 三个正式阶段\n"
    "• 记忆管理：维护会话历史作为短期记忆，支持 Agent 回溯之前的决策\n"
    "• 实验表明：ChatDev 能在 7 分钟内以不到 1 美元的成本完成一个小型软件项目"
)

doc.add_heading("3.2 AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation", 2)
doc.add_paragraph("作者：Qingyun Wu et al. (Microsoft Research, 2023)")
doc.add_paragraph(
    "AutoGen 是一个多 Agent 对话编排框架，核心抽象是「可对话智能体（Conversable Agent）」——"
    "任何能接收消息并产生回复的实体。通过组合不同类型的 Agent（AssistantAgent、UserProxyAgent、GroupChatManager），"
    "开发者可以灵活构建复杂的多 Agent 协作工作流。\n\n"
    "框架设计：\n"
    "• 对话驱动：Agent 之间的所有交互均通过消息传递完成，统一抽象\n"
    "• 人机回环：UserProxyAgent 可在关键决策点暂停等待人类输入\n"
    "• 工具集成：Agent 可注册任意 Python 函数作为工具供 LLM 调用\n"
    "• 群聊模式：GroupChat 中多个 Agent 轮流发言，由 Manager 决定发言顺序\n"
    "• 在代码生成、数学推理、数据分析等场景实现了显著优于单 Agent 的效果"
)

doc.add_heading("3.3 AgentVerse: A Multi-Agent Framework for Collaborative Problem Solving", 2)
doc.add_paragraph("作者：Weize Chen et al. (Tsinghua University, 2023)")
doc.add_paragraph(
    "AgentVerse 提出了一个多 Agent 协作解决复杂问题的通用框架，包含四个核心阶段："
    "Expert Recruitment（根据问题招募最合适的专家 Agent）、Collaborative Decision-Making（多轮讨论达成一致）、"
    "Action Execution（按照决策分配并执行子任务）、Evaluation（对执行结果进行多维度评估）。"
    "该框架的核心贡献在于将 Agent 协作过程形式化为可优化的流程，并证明结构化的多 Agent 讨论比单体 Agent 的 Chain-of-Thought 更有优势。"
)

# ── 四、RL 训练方法与基础架构 ──
doc.add_heading("四、RL 训练方法与基础架构", 1)

doc.add_heading("4.1 GRPO: Group Relative Policy Optimization", 2)
doc.add_paragraph("作者：DeepSeek-AI (2024)")
doc.add_paragraph(
    "GRPO 是 DeepSeek 团队在 PPO 基础上改进的策略优化算法，专为 LLM 的 RL 训练设计。"
    "其核心创新在于取消了传统 PPO 中所需的 Critic 模型（价值网络），转而通过「组内相对比较」来计算优势函数。\n\n"
    "算法原理：\n"
    "• 对每个 prompt 采样 G 个输出，用组内均值作为 baseline 计算相对优势\n"
    "• 优势函数：A_i = (r_i - mean(r)) / std(r)，其中 r_i 是第 i 个输出的奖励\n"
    "• 省去了与策略模型同等规模的价值网络，显存占用几乎减半\n"
    "• 加入 KL 散度惩罚项（直接加在损失函数中）以防止策略更新过于激进\n"
    "• DeepSeek-R1、DeepSeekMath 等工作均基于 GRPO 训练"
)

doc.add_heading("4.2 RLHF vs RLAIF: Scaling Alignment with AI Feedback", 2)
doc.add_paragraph("作者：Anthropic (2023, Constitutional AI)")
doc.add_paragraph(
    "RLHF（Reinforcement Learning from Human Feedback）是当前大模型对齐的事实标准方法，"
    "但其受限于人类标注的成本与一致性。RLAIF（RL from AI Feedback）探索用 AI 自身的判断替代人类偏好标注，"
    "实现可扩展的对齐训练。Anthropic 提出的 Constitutional AI 是 RLAIF 的代表性方法。\n\n"
    "Constitutional AI 流程：\n"
    "1. 监督学习阶段：用宪法（Constitution，即一组行为准则）指导模型自我修订有害输出\n"
    "2. RL 阶段：基于宪法反馈训练的偏好模型替代人工标注的奖励模型\n"
    "3. 实验显示：Constitutional AI 训练的模型在有用性上持平 RLHF，在无害性上超过 RLHF\n"
    "4. 关键优势：宪法可被审计和修订，使对齐过程更具透明性和可控性"
)

doc.add_heading("4.3 ReAct: Synergizing Reasoning and Acting in Language Models", 2)
doc.add_paragraph("作者：Shunyu Yao et al. (2023)")
doc.add_paragraph(
    "ReAct 是推理-行动交错范式的奠基性工作。它提出将 Chain-of-Thought 推理与 Tool-Use 行动交替进行——"
    "模型先思考（Thought）、再行动（Action）、然后观察结果（Observation），如此循环直至任务完成。"
    "这一范式已成为几乎所有现代 Agent 框架的标准架构。\n\n"
    "实验发现：\n"
    "• 单独的 Chain-of-Thought 容易产生幻觉（无法验证推理正确性）\n"
    "• 单独的 Action-only 缺乏推理规划能力（无法应对需要多步推理的复杂任务）\n"
    "• ReAct 结合两者：推理指导行动，行动的反馈修正推理\n"
    "• 在 HotpotQA 和 Fever 等知识密集型任务上，ReAct 显著优于纯推理或纯行动的 baseline"
)

# ── 五、总结与展望 ──
doc.add_heading("五、总结与展望", 1)
doc.add_paragraph(
    "从上述论文可以看出，Agent RL 领域正沿着以下几条主线快速发展：\n\n"
    "1. 从稀疏奖励到密集奖励：DeepSeek-R1 展示了仅用最终奖励就能涌现推理能力的惊人结果，"
    "但 STILL-ALIVE 等工作在进一步探索如何以低成本提供步级奖励信号，两者的结合是未来方向。\n\n"
    "2. 从单 Agent 到多 Agent：ChatDev、AutoGen 等工作证明了多 Agent 协作在复杂任务上的优势，"
    "但如何用 RL 优化 Agent 间的协作策略仍是开放问题。\n\n"
    "3. 从静态到自我进化：DeepEnlighten 提出的循环训练范式让 Agent 从使用中学习，"
    "这种「部署即训练」的思路有望成为构建持续进化 AI 系统的关键。\n\n"
    "4. 从 API 依赖到本地部署：BGE、Qwen 等开源模型的快速发展使得高质量 Embedding 和推理能力可以本地运行，"
    "为 RAG 系统（如本项目 RAGent）提供了坚实的技术基础。\n\n"
    "参考文献：本文档涉及的所有论文均可在 arXiv 上通过标题检索到全文。"
)

import os
os.makedirs("data/raws/测试空间", exist_ok=True)
doc.save("data/raws/测试空间/Agent_RL_Papers.docx")
print("文档已生成: data/raws/测试空间/Agent_RL_Papers.docx")
